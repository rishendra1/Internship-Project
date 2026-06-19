import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set internal cell padding (margins) in dxa (1/20th of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Apply professional thin gray borders to the table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="808080"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:insideV w:val="none"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def add_paragraph_with_runs(doc, text, style=None, space_after=6, line_spacing=1.15):
    """Adds a paragraph and parses simple **bold** and *italic* markdown formatting."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    # Simple regex tokenizer for markdown bold and italic
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = p.add_run(token[1:-1])
            run.italic = True
        else:
            p.add_run(token)
    return p

def convert_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Configure page settings (Standard 1 inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styling configuration
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Colors
    color_primary = RGBColor(0x1B, 0x36, 0x5D)  # Deep Navy Blue
    color_secondary = RGBColor(0x5C, 0x76, 0x8D) # Slate Blue
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    in_code_block = False
    code_lines = []
    
    in_mermaid = False
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                # Check if it is a mermaid block
                if 'mermaid' in line:
                    in_mermaid = True
                else:
                    in_mermaid = False
                code_lines = []
            else:
                in_code_block = False
                if not in_mermaid:
                    # Write code block to Word
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.space_before = Pt(6)
                    
                    # Style code container
                    # Add background shading to the paragraph (requires XML manipulation)
                    pPr = p._p.get_or_add_pPr()
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                    pPr.append(shd)
                    
                    # Add border
                    pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="1B365D"/></w:pBdr>')
                    pPr.append(pbdr)
                    
                    code_text = "\n".join(code_lines)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue
            
        if in_code_block:
            if not in_mermaid:
                code_lines.append(line)
            i += 1
            continue
            
        # Handle Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(22)
            run.bold = True
            run.font.color.rgb = color_primary
            i += 1
            continue
            
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = color_primary
            i += 1
            continue
            
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = color_secondary
            i += 1
            continue
            
        # Handle horizontal rules
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            pPr = p._p.get_or_add_pPr()
            pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D3D3D3"/></w:pBdr>')
            pPr.append(pbdr)
            i += 1
            continue
            
        # Handle bullet points
        if line.strip().startswith('* ') or line.strip().startswith('- '):
            clean_line = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            # Formatting bullet text (handle bold/italic)
            tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', clean_line)
            for token in tokens:
                if token.startswith('**') and token.endswith('**'):
                    run = p.add_run(token[2:-2])
                    run.bold = True
                elif token.startswith('*') and token.endswith('*'):
                    run = p.add_run(token[1:-1])
                    run.italic = True
                else:
                    p.add_run(token)
            i += 1
            continue
            
        # Handle numbered lists
        match_num = re.match(r'^(\d+)\.\s(.*)', line.strip())
        if match_num:
            num = match_num.group(1)
            text = match_num.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for token in tokens:
                if token.startswith('**') and token.endswith('**'):
                    run = p.add_run(token[2:-2])
                    run.bold = True
                elif token.startswith('*') and token.endswith('*'):
                    run = p.add_run(token[1:-1])
                    run.italic = True
                else:
                    p.add_run(token)
            i += 1
            continue
            
        # Handle Markdown Tables
        if line.strip().startswith('|'):
            in_table = True
            # Read all consecutive table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            # Process table
            if len(table_lines) >= 3:
                # Parse headers
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                
                # Check column alignments from separator (row index 1)
                aligns = []
                sep_cells = table_lines[1].split('|')[1:-1]
                for cell in sep_cells:
                    c = cell.strip()
                    if c.startswith(':') and c.endswith(':'):
                        aligns.append(WD_ALIGN_PARAGRAPH.CENTER)
                    elif c.endswith(':'):
                        aligns.append(WD_ALIGN_PARAGRAPH.RIGHT)
                    else:
                        aligns.append(WD_ALIGN_PARAGRAPH.LEFT)
                        
                # Parse rows
                rows_data = []
                for r_line in table_lines[2:]:
                    row_cells = [cell.strip() for cell in r_line.split('|')[1:-1]]
                    # Ensure matching length
                    while len(row_cells) < len(headers):
                        row_cells.append("")
                    rows_data.append(row_cells[:len(headers)])
                    
                # Create Word table
                table = doc.add_table(rows=1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                # Format headers
                hdr_cells = table.rows[0].cells
                for col_idx, text in enumerate(headers):
                    cell = hdr_cells[col_idx]
                    set_cell_background(cell, "1B365D")  # Deep Navy Blue
                    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
                    
                    p = cell.paragraphs[0]
                    p.alignment = aligns[col_idx]
                    # Parse bold markers in headers
                    clean_text = text.replace("**", "").replace("*", "")
                    run = p.add_run(clean_text)
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White text
                    
                # Add data rows
                for row_idx, r_data in enumerate(rows_data):
                    row = table.add_row()
                    # Zebra striping: alternate background colors
                    bg_color = "F9FBFD" if row_idx % 2 == 0 else "FFFFFF"
                    
                    for col_idx, cell_value in enumerate(r_data):
                        cell = row.cells[col_idx]
                        set_cell_background(cell, bg_color)
                        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                        
                        p = cell.paragraphs[0]
                        p.alignment = aligns[col_idx]
                        p.paragraph_format.space_after = Pt(0)
                        
                        # Add value with rich runs
                        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', cell_value)
                        for token in tokens:
                            if token.startswith('**') and token.endswith('**'):
                                r_run = p.add_run(token[2:-2])
                                r_run.bold = True
                            elif token.startswith('*') and token.endswith('*'):
                                r_run = p.add_run(token[1:-1])
                                r_run.italic = True
                            else:
                                r_run = p.add_run(token)
                            r_run.font.name = 'Calibri'
                            r_run.font.size = Pt(9.5)
                            
                # Add spacing after table
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_after = Pt(12)
                p_spacer.paragraph_format.space_before = Pt(6)
            continue
            
        # Handle regular text
        if line.strip() != "":
            # Skip empty lines or format paragraphs
            add_paragraph_with_runs(doc, line, space_after=6)
            
        i += 1
        
    doc.save(docx_path)
    print(f"Success: Saved docx report to {docx_path}")

if __name__ == '__main__':
    md_file = r"c:\Users\rishe\PycharmProjects\PythonProject\final_project_report.md"
    docx_file = r"c:\Users\rishe\PycharmProjects\PythonProject\final_project_report.docx"
    convert_markdown_to_docx(md_file, docx_file)
